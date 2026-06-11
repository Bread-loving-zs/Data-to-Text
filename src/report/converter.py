import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from src.config import setup_logging

logger = setup_logging(__name__)


class MarkdownToDocx:
    def __init__(self):
        self.doc = Document()

        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "宋体"
        font.size = Pt(12)

        for level in range(1, 4):
            heading_style = self.doc.styles[f"Heading {level}"]
            heading_style.font.color.rgb = RGBColor(0x1A, 0x52, 0x7E)
            if level == 1:
                heading_style.font.size = Pt(22)
            elif level == 2:
                heading_style.font.size = Pt(16)
            else:
                heading_style.font.size = Pt(14)

        self.default_sections = self.doc.sections
        for section in self.default_sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.18)
            section.right_margin = Cm(3.18)

    def convert(self, markdown_text: str) -> Document:
        lines = markdown_text.strip().split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            if not line.strip():
                i += 1
                continue

            if line.strip().startswith("```"):
                i = self._parse_code_block(lines, i)
                continue

            stripped = line.strip()

            if re.match(r"^(-{3,}|\*{3,}|_{3,})$", stripped):
                self.doc.add_paragraph("—" * 40)
                i += 1
                continue

            if stripped.startswith("> "):
                i = self._parse_blockquote(lines, i)
                continue

            heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                self.doc.add_heading(text, level=min(level, 3))
                i += 1
                continue

            table_match = re.match(r"^\|(.+)\|$", line)
            if table_match:
                i = self._parse_table(lines, i)
                continue

            image_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line)
            if image_match:
                alt_text = image_match.group(1)
                img_path = image_match.group(2)
                self._add_image(img_path, alt_text)
                i += 1
                continue

            unordered_match = re.match(r"^(\s*)([-*+])\s+(.+)$", line)
            if unordered_match:
                i = self._parse_unordered_list(lines, i)
                continue

            if self._is_ordered_list_item(line):
                i = self._parse_ordered_list(lines, i)
                continue

            processed = self._process_inline(line)
            self.doc.add_paragraph(processed)
            i += 1

        return self.doc

    def _is_ordered_list_item(self, line: str) -> bool:
        return bool(re.match(r"^\s*\d+\.\s+.+", line))

    def _process_inline(self, text: str) -> str:
        text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
        text = re.sub(r"\*\*(.+?)\*\*", r"【\1】", text)
        return text

    def _parse_unordered_list(self, lines: list[str], start_idx: int) -> int:
        i = start_idx
        while i < len(lines):
            line = lines[i]
            match = re.match(r"^(\s*)([-*+])\s+(.+)$", line)
            if not match:
                break
            indent = len(match.group(1))
            text = match.group(3)
            sub_match = re.match(r"^([-*+])\s+(.+)$", text)
            if sub_match:
                level = indent // 2 + 1
                text = sub_match.group(2)
            else:
                level = indent // 2
            processed = self._process_inline(text)
            if level <= 1:
                self.doc.add_paragraph(processed, style='List Bullet')
            elif level == 2:
                self.doc.add_paragraph(processed, style='List Bullet 2')
            else:
                self.doc.add_paragraph(processed, style='List Bullet 3')
            i += 1
        return i

    def _parse_ordered_list(self, lines: list[str], start_idx: int) -> int:
        i = start_idx
        while i < len(lines):
            if not self._is_ordered_list_item(lines[i]):
                break
            match = re.match(r"^\s*\d+\.\s+(.+)$", lines[i])
            text = self._process_inline(match.group(1))
            self.doc.add_paragraph(text, style='List Number')
            i += 1
        return i

    def _parse_blockquote(self, lines: list[str], start_idx: int) -> int:
        i = start_idx
        quote_lines = []
        while i < len(lines):
            stripped = lines[i].strip()
            if stripped.startswith("> "):
                quote_lines.append(stripped[2:])
                i += 1
            elif stripped.startswith(">") and stripped != ">":
                quote_lines.append(stripped[1:].strip())
                i += 1
            else:
                break
        text = " ".join(quote_lines)
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(self._process_inline(text))
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.italic = True
        return i

    def _parse_code_block(self, lines: list[str], start_idx: int) -> int:
        i = start_idx + 1
        code_lines = []
        while i < len(lines):
            if lines[i].strip().startswith("```"):
                i += 1
                break
            code_lines.append(lines[i])
            i += 1
        for code_line in code_lines:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(code_line if code_line else " ")
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        return i

    def _parse_table(self, lines: list[str], start_idx: int) -> int:
        table_lines = []
        i = start_idx
        while i < len(lines) and lines[i].strip().startswith("|"):
            table_lines.append(lines[i].strip())
            i += 1

        if len(table_lines) < 2:
            return i

        header_cells = [c.strip() for c in table_lines[0].split("|")[1:-1]]
        data_rows = []
        for line in table_lines[1:]:
            if re.match(r"^[\|\s\-:]+$", line):
                continue
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if cells:
                data_rows.append(cells)

        if not header_cells:
            return i

        num_cols = len(header_cells)
        table = self.doc.add_table(rows=1 + len(data_rows), cols=num_cols)
        table.style = "Light Grid Accent 1"

        for col_idx, cell_text in enumerate(header_cells):
            cell = table.rows[0].cells[col_idx]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        for row_idx, row_data in enumerate(data_rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.rows[row_idx + 1].cells[col_idx]
                    cell.text = cell_text

        return i

    def _add_image(self, img_path: str, alt_text: str = ""):
        if Path(img_path).exists():
            self.doc.add_picture(img_path, width=Inches(5.5))
            last_paragraph = self.doc.paragraphs[-1] if self.doc.paragraphs else self.doc.add_paragraph()
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if alt_text:
                caption = self.doc.add_paragraph(f"图：{alt_text}")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.style = self.doc.styles["Normal"]
        else:
            self.doc.add_paragraph(f"[图片未找到: {img_path}]")

    def save(self, filepath: Path):
        filepath.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(filepath))
        return filepath


def markdown_to_docx(markdown_text: str, output_path: Path) -> Path:
    converter = MarkdownToDocx()
    converter.convert(markdown_text)
    return converter.save(output_path)